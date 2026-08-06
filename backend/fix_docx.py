import sys
from docx import Document

def main():
    template_path = "/home/mahdi/CascadeProjects/ndc/backend/data/Template NDC_modified.docx"
    doc = Document(template_path)
    
    # We want to find {{ photo_img }} and remove it.
    # And find "PHOTO DE SITE" and replace it with {{ photo_img }}.
    
    # 1. Look through all paragraphs and tables
    for p in doc.paragraphs:
        if '{{ photo_img }}' in p.text:
            p.text = p.text.replace('{{ photo_img }}', '')
        if '{{photo_img}}' in p.text:
            p.text = p.text.replace('{{photo_img}}', '')
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '{{ photo_img }}' in p.text:
                        p.text = p.text.replace('{{ photo_img }}', '')
                    if '{{photo_img}}' in p.text:
                        p.text = p.text.replace('{{photo_img}}', '')
                    
                    if 'PHOTO DE SITE' in p.text:
                        p.text = p.text.replace('PHOTO DE SITE', '{{ photo_img }}')
                        
    # Ensure docxtpl variables are correctly formatted
    doc.save(template_path)
    print("Done")

if __name__ == "__main__":
    main()
